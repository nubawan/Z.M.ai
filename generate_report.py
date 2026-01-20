"""Generate CCP Report for Z.M.ai"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('CCP Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_heading('RAG-Based Academic Policy Chatbot', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Submitted by section
info = doc.add_paragraph()
info.add_run('Submitted by: Zayan Shahid & Maviya').bold = True
info.add_run('\nCourse: Parallel & Distributed Computing')
info.add_run('\nSubmitted to: Dr Ali Akbar Siddique')
info.add_run('\nDay: Friday/Saturday')
info.add_run('\nTime Slot: 8:30-9:50')

# Add spacing
doc.add_paragraph()

# 1. Introduction
doc.add_heading('1. Introduction', 2)
p = doc.add_paragraph(
    'Z.M.ai (ZayMayAi) is an AI-powered academic policy assistant chatbot developed using the '
    'Retrieval-Augmented Generation (RAG) approach. Unlike traditional chatbots that rely solely '
    'on pre-trained language models, Z.M.ai enhances response accuracy by retrieving relevant '
    'information from academic policy documents and web content before generating answers. '
    'The system is deployed as an interactive web application using Streamlit, enabling '
    'real-time, user-friendly interaction for students seeking information about university policies.'
)
doc.add_paragraph('Key highlights:', style='List Bullet')
doc.add_paragraph('Combines web scraping and PDF content retrieval with natural language generation.')
doc.add_paragraph('Provides context-aware and domain-specific responses for academic policies.')
doc.add_paragraph('Offers a clean, professional, and interactive web-based chat interface.')
doc.add_paragraph('Uses keyword-based retrieval for efficient and accurate responses.')

# 2. Background
doc.add_heading('2. Background', 2)
p = doc.add_paragraph(
    'Recent advancements in Large Language Models (LLMs) have significantly improved conversational '
    'AI systems. However, standalone LLMs often suffer from limitations such as hallucinations, '
    'lack of domain specificity, and inability to reference updated or external knowledge. '
    'Retrieval-Augmented Generation addresses these issues by integrating a retrieval mechanism '
    'that fetches relevant documents and supplies them as context to the LLM.'
)
doc.add_paragraph('Benefits of RAG:', style='List Bullet')
doc.add_paragraph('Improved factual accuracy for academic policy queries.')
doc.add_paragraph('Reduced hallucinations in domain-specific responses.')
doc.add_paragraph('Better performance on policy-related questions.')
doc.add_paragraph('Scalable knowledge updates without retraining the model.')

# 3. Problem Statement
doc.add_heading('3. Problem Statement', 2)
doc.add_paragraph('Students often face challenges when searching for accurate academic policy information:')
doc.add_paragraph('University policies are scattered across multiple documents and web pages.', style='List Bullet')
doc.add_paragraph('Manual searching is time-consuming and error-prone.', style='List Bullet')
doc.add_paragraph('Traditional chatbots provide generic or vague responses to policy queries.', style='List Bullet')
doc.add_paragraph('Lack of a centralized, intelligent system for academic policy information.', style='List Bullet')

p = doc.add_paragraph('Proposed Solution:')
p.add_run(
    '\nTo overcome these issues, Z.M.ai implements a RAG-based architecture that retrieves '
    'relevant information from academic policy documents and web content, then uses an LLM '
    'to generate accurate, context-aware responses for student queries.'
)

# 4. Objectives
doc.add_heading('4. Objectives', 2)
doc.add_paragraph('The primary objectives of this project are:')
doc.add_paragraph('To design and implement a RAG-based academic policy assistant chatbot.', style='List Number')
doc.add_paragraph('To integrate web scraping and PDF processing for comprehensive knowledge base.', style='List Number')
doc.add_paragraph('To develop a professional web-based chat interface using Streamlit.', style='List Number')
doc.add_paragraph('To demonstrate accurate, context-aware AI responses for academic policy queries.', style='List Number')
doc.add_paragraph('To provide a simple, efficient solution without complex dependencies.', style='List Number')

# 5. Scope
doc.add_heading('5. Scope of the Project', 2)
p = doc.add_paragraph(
    'The scope of the Z.M.ai project defines the functional boundaries and capabilities of the '
    'system implemented during this project. The chatbot is designed to focus on accuracy, '
    'simplicity, and real-time interaction while demonstrating the core concepts of '
    'Retrieval-Augmented Generation.'
)
doc.add_paragraph('The scope of Z.M.ai includes:', style='List Bullet')
doc.add_paragraph('Text-based user queries for academic policies in natural language.')
doc.add_paragraph('Web scraping from official university policy pages.')
doc.add_paragraph('PDF document processing for academic policy manuals.')
doc.add_paragraph('Keyword-based retrieval for efficient information matching.')
doc.add_paragraph('Real-time response generation using Groq LLM.')
doc.add_paragraph('Professional, user-friendly web interface using Streamlit.')

p = doc.add_paragraph('Out of Scope (Future Enhancements):')
p.add_run('\n').bold = True
doc.add_paragraph('Voice-based queries and responses.', style='List Bullet')
doc.add_paragraph('Multi-language support beyond English.', style='List Bullet')
doc.add_paragraph('User authentication and personalized conversation history.', style='List Bullet')
doc.add_paragraph('Integration with university student portal systems.', style='List Bullet')

# 6. Tools and Technologies
doc.add_heading('6. Tools and Technologies Required', 2)
p = doc.add_paragraph(
    'The Z.M.ai project utilizes modern AI and web development tools to ensure efficiency, '
    'simplicity, and ease of deployment:'
)
doc.add_paragraph('Programming Language: ', style='List Bullet')
p.add_run('Python - Backend logic, data processing, and AI integration')
doc.add_paragraph('Frontend Framework: ', style='List Bullet')
p.add_run('Streamlit - Lightweight and interactive web interface')
doc.add_paragraph('LLM Provider: ', style='List Bullet')
p.add_run('Groq API - Fast, cost-effective Llama 3.1 models for natural language responses')
doc.add_paragraph('Retrieval Method: ', style='List Bullet')
p.add_run('Keyword-based scoring - Efficient matching without complex vector databases')
doc.add_paragraph('Web Scraping: ', style='List Bullet')
p.add_run('BeautifulSoup & Requests - HTML parsing and content extraction')
doc.add_paragraph('PDF Processing: ', style='List Bullet')
p.add_run('pdfplumber - Text extraction from PDF documents')
doc.add_paragraph('Deployment Platform: ', style='List Bullet')
p.add_run('Streamlit Cloud - Easy and accessible web deployment')

# 7. System Design
doc.add_heading('7. System Design and Architecture', 2)
doc.add_heading('7.1 Architecture Overview', 3)
p = doc.add_paragraph(
    'Z.M.ai follows a simplified RAG architecture, ensuring efficiency and ease of maintenance. '
    'Each module performs a specific function within the system:'
)
doc.add_paragraph('User Interface Layer: ', style='List Bullet')
p.add_run('Provides a Streamlit-based chat interface with professional styling')
doc.add_paragraph('Data Collection Layer: ', style='List Bullet')
p.add_run('Scraps web content from university policy pages and extracts text from PDF documents')
doc.add_paragraph('Caching Layer: ', style='List Bullet')
p.add_run('Stores scraped content locally for fast loading and reduced API calls')
doc.add_paragraph('Retrieval Layer: ', style='List Bullet')
p.add_run('Uses keyword-based scoring to find relevant content chunks based on user query terms')
doc.add_paragraph('LLM Generation Layer: ', style='List Bullet')
p.add_run('Uses Groqs Llama 3.1 model to generate accurate, context-aware responses')
doc.add_paragraph('Display Layer: ', style='List Bullet')
p.add_run('Presents generated responses in the Streamlit chat interface with message history')

doc.add_heading('7.2 Flow Diagram (Logical Flow)', 3)

# 8. Methodology
doc.add_heading('8. Methodology / Implementation Steps', 2)
p = doc.add_paragraph('The project was implemented using the following systematic steps:')
doc.add_paragraph('Content Collection: Scraping university policy website and processing PDF documents.', style='List Number')
doc.add_paragraph('Local Storage: Caching content locally for faster access and reduced dependencies.', style='List Number')
doc.add_paragraph('Keyword Processing: Extracting meaningful terms from user queries.', style='List Number')
doc.add_paragraph('Context Retrieval: Scoring and retrieving relevant text chunks using keyword matching.', style='List Number')
doc.add_paragraph('Related Terms: Adding scores for related terms (e.g., attendance to present/absent, grade to GPA/CGPA).', style='List Number')
doc.add_paragraph('Answer Generation: Passing retrieved context and query to Groq LLM.', style='List Number')
doc.add_paragraph('UI Integration: Displaying conversation with professional styling using Streamlit.', style='List Number')
doc.add_paragraph('Testing: Validating accuracy, relevance, and performance with various queries.', style='List Number')

# 9. Languages and Libraries
doc.add_heading('9. Languages and Libraries Used', 2)
doc.add_paragraph('Python: ', style='List Bullet')
p.add_run('Core backend logic (~380 lines single-file architecture)')
doc.add_paragraph('HTML/CSS: ', style='List Bullet')
p.add_run('Professional UI styling within Streamlit markdown')
doc.add_paragraph('\nLibraries:', style='List Bullet')
doc.add_paragraph('streamlit', style='List Bullet 2')
doc.add_paragraph('groq', style='List Bullet 2')
doc.add_paragraph('requests', style='List Bullet 2')
doc.add_paragraph('beautifulsoup4', style='List Bullet 2')
doc.add_paragraph('pdfplumber', style='List Bullet 2')

# 10. APIs
doc.add_heading('10. APIs Used and Integration', 2)
doc.add_heading('10.1 Groq LLM API', 3)
p = doc.add_paragraph('Purpose: Generate natural language responses using retrieved context.')
p = doc.add_paragraph('\nIntegration Process:')
doc.add_paragraph('Install Groq SDK.', style='List Number')
doc.add_paragraph('Configure API key via Streamlit secrets.', style='List Number')
doc.add_paragraph('Send user query along with retrieved context.', style='List Number')
doc.add_paragraph('Receive and display generated response.', style='List Number')

doc.add_heading('10.2 Web Scraping', 3)
doc.add_paragraph('Purpose: Extract policy content from university website.')
doc.add_paragraph('Process:', style='List Bullet')
doc.add_paragraph('HTTP requests with proper headers')
doc.add_paragraph('HTML parsing with BeautifulSoup')
doc.add_paragraph('Content cleaning and formatting')

# 11. Code
doc.add_heading('11. Code Implementation (Sample)', 2)
code_text = '''# Main application structure
def main():
    st.set_page_config(page_title="Z.M.ai | Academic Policy Assistant")

    # Load knowledge base (cached locally)
    if not st.session_state.knowledge_loaded:
        st.session_state.knowledge_base = load_or_create_knowledge_base()

    # Chat interface
    if prompt := st.chat_input("Ask about academic policies..."):
        # Retrieve relevant context
        context = retrieve_context(prompt, st.session_state.knowledge_base)

        # Generate response
        response = get_llm_response(prompt, context, history)

        # Display response
        st.markdown(response)'''

p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Courier New'

# 12. Challenges
doc.add_heading('12. Challenges Faced', 2)
doc.add_paragraph('Balancing simplicity with accuracy in retrieval mechanism.', style='List Bullet')
doc.add_paragraph('Handling diverse document formats (PDF, web content).', style='List Bullet')
doc.add_paragraph('Optimizing keyword matching for relevant results.', style='List Bullet')
doc.add_paragraph('Managing Streamlit session state for chat history.', style='List Bullet')
doc.add_paragraph('Creating professional UI within Streamlit constraints.', style='List Bullet')

# 13. Results
doc.add_heading('13. Results and Outcomes', 2)
doc.add_paragraph('Successfully implemented a functional RAG-based academic policy assistant.', style='List Bullet')
doc.add_paragraph('Achieved accurate and context-aware responses for policy queries.', style='List Bullet')
doc.add_paragraph('Deployed working web application accessible via browser.', style='List Bullet')
doc.add_paragraph('Implemented efficient caching system for fast knowledge base loading.', style='List Bullet')
doc.add_paragraph('Created professional, polished user interface.', style='List Bullet')

# 14. Conclusion
doc.add_heading('14. Conclusion', 2)
p = doc.add_paragraph(
    'Z.M.ai successfully demonstrates the practical implementation of a Retrieval-Augmented '
    'Generation based chatbot tailored for academic policy assistance. By integrating web scraping, '
    'PDF processing, and intelligent keyword-based retrieval with a Large Language Model, the '
    'system is able to generate accurate, relevant, and context-aware responses for student queries. '
    'The simplified architecture eliminates complex dependencies while maintaining high response quality. '
    'The use of Groqs fast Llama models ensures quick response times. The Streamlit web interface '
    'provides a professional and user-friendly interaction experience. This project highlights how '
    'RAG enhances the performance of AI chatbots in domain-specific scenarios. It also showcases the '
    'effective integration of modern AI tools within a clean, maintainable architecture. Overall, '
    'Z.M.ai fulfills its intended objectives successfully and provides valuable assistance to students '
    'seeking academic policy information.'
)

# 15. Future Improvements
doc.add_heading('15. Future Improvements', 2)
doc.add_paragraph('Support for multiple PDF documents and policy sources.', style='List Bullet')
doc.add_paragraph('Persistent conversation memory across sessions.', style='List Bullet')
doc.add_paragraph('Advanced semantic search with sentence transformers.', style='List Bullet')
doc.add_paragraph('Multi-lingual query support for diverse student populations.', style='List Bullet')
doc.add_paragraph('Voice input/output capabilities.', style='List Bullet')
doc.add_paragraph('Integration with university database for real-time policy updates.', style='List Bullet')

# 16. References
doc.add_heading('16. References', 2)
doc.add_paragraph('Groq API Documentation - https://console.groq.com/docs', style='List Bullet')
doc.add_paragraph('Streamlit Official Documentation.', style='List Bullet')
doc.add_paragraph('GitHub Repository: https://github.com/nubawan/Z.M.ai', style='List Bullet')
doc.add_paragraph('Deployed Web App: https://zaymavai.streamlit.app/', style='List Bullet')

# 17. Output
doc.add_heading('17. Output', 2)
p = doc.add_paragraph(
    'The Z.M.ai chatbot is successfully deployed and accessible at:\n'
    'https://zaymavai.streamlit.app/\n\n'
    'The application features:'
)
doc.add_paragraph('Professional gradient UI with custom styling.', style='List Bullet')
doc.add_paragraph('Real-time chat interface with message history.', style='List Bullet')
doc.add_paragraph('Fast knowledge base loading from local cache.', style='List Bullet')
doc.add_paragraph('Accurate policy information retrieval.', style='List Bullet')
doc.add_paragraph('Builder credits (Zayan & Maviya) displayed in header and footer.', style='List Bullet')

# Save
doc.save('ZMai_CCP_Report.docx')
print('Report generated successfully: ZMai_CCP_Report.docx')
